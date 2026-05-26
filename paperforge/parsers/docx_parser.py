"""Parse Word .docx into a Paper, using built-in heading styles when present."""

from __future__ import annotations

from pathlib import Path

import docx  # python-docx

from ..model import Paper, Table
from .heuristics import Block, build_paper


def _style_level(style_name: str):
    """Map a Word paragraph style to a heading level, or None for body text."""
    if not style_name:
        return None
    name = style_name.lower()
    if name == "title":
        return None  # handled separately as the document title
    if name.startswith("heading"):
        tail = name.replace("heading", "").strip()
        try:
            return int(tail)
        except ValueError:
            return 1
    return None


def parse_docx(path: str) -> Paper:
    document = docx.Document(path)
    blocks = []
    title_from_style = None

    for para in document.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        style = para.style.name if para.style else ""
        if style and style.lower() == "title" and title_from_style is None:
            title_from_style = text
            continue
        level = _style_level(style)
        blocks.append(Block(text=text, level=level if level else 0))

    # If Word styles gave us no headings at all, let the heuristics guess by
    # resetting known body levels to "unknown".
    if not any(b.level for b in blocks):
        blocks = [Block(text=b.text, level=None) for b in blocks]

    # Prepend an explicit Title block so build_paper picks it up first.
    if title_from_style:
        blocks.insert(0, Block(text=title_from_style, level=0))

    paper = build_paper(blocks)
    if title_from_style:
        paper.title = title_from_style

    # Capture tables verbatim (best-effort; renderers may use them).
    for tbl in document.tables:
        rows = [[cell.text.strip() for cell in row.cells] for row in tbl.rows]
        if rows:
            paper.tables.append(Table(rows=rows))

    if not paper.title:
        paper.title = Path(path).stem
    return paper
