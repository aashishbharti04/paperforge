"""Render a Paper to a formatted Word .docx for IEEE, Springer, or Elsevier.

We build the formatting programmatically (fonts, sizes, columns, spacing) to
approximate each publisher's house style. The result is an editable, submission-
ready starting point — final camera-ready tweaks should still be checked
against the publisher's official template.
"""

from __future__ import annotations

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.shared import Pt, Inches, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from ..model import Paper, Section


# --- low-level helpers ------------------------------------------------------

def _set_columns(section, num: int, space_twips: int = 360):
    """Set the number of newspaper-style columns on a section."""
    sectPr = section._sectPr
    cols = sectPr.find(qn("w:cols"))
    if cols is None:
        cols = OxmlElement("w:cols")
        sectPr.append(cols)
    cols.set(qn("w:num"), str(num))
    cols.set(qn("w:space"), str(space_twips))


def _base_font(document, name: str, size: int):
    style = document.styles["Normal"]
    style.font.name = name
    style.font.size = Pt(size)
    # Ensure east-asian/complex scripts also use the font.
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    for attr in ("w:ascii", "w:hAnsi", "w:cs"):
        rfonts.set(qn(attr), name)


def _para(document, text, *, size=None, bold=False, italic=False,
          align=None, space_after=6, font=None, color=None, caps=False):
    p = document.add_paragraph()
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(0)
    run = p.add_run(text.upper() if caps else text)
    run.bold = bold
    run.italic = italic
    if size:
        run.font.size = Pt(size)
    if font:
        run.font.name = font
    if color:
        run.font.color.rgb = color
    return p


# --- shared section writer --------------------------------------------------

def _write_sections(document, sections, body_size, numbering, prefix=""):
    """Write sections recursively. ``numbering`` controls the heading scheme."""
    for i, sec in enumerate(sections, 1):
        num = f"{prefix}{i}"
        if sec.level == 1:
            heading = _heading_text(num, sec.title, numbering)
            _para(document, heading, size=body_size + 1, bold=True,
                  space_after=4)
        else:
            heading = _heading_text(num, sec.title, numbering, sub=True)
            _para(document, heading, size=body_size, bold=True, italic=True,
                  space_after=3)
        for para in sec.paragraphs:
            p = _para(document, para, size=body_size,
                      align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=6)
            p.paragraph_format.first_line_indent = Inches(0.2)
        if sec.children:
            _write_sections(document, sec.children, body_size, numbering,
                            prefix=f"{num}.")


def _heading_text(num, title, numbering, sub=False):
    if numbering == "roman" and not sub:
        return f"{_to_roman(num)}.  {title.upper()}"
    if numbering == "roman" and sub:
        return f"{num}  {title}"
    return f"{num}  {title}"


def _to_roman(num_str):
    try:
        n = int(str(num_str).split(".")[0])
    except ValueError:
        return num_str
    vals = [(1000, "M"), (900, "CM"), (500, "D"), (400, "CD"), (100, "C"),
            (90, "XC"), (50, "L"), (40, "XL"), (10, "X"), (9, "IX"),
            (5, "V"), (4, "IV"), (1, "I")]
    out = ""
    for v, sym in vals:
        while n >= v:
            out += sym
            n -= v
    return out


def _write_references(document, paper, size):
    if not paper.references:
        return
    _para(document, "REFERENCES", size=size + 1, bold=True, space_after=4)
    for i, ref in enumerate(paper.references, 1):
        p = document.add_paragraph()
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.left_indent = Inches(0.25)
        p.paragraph_format.first_line_indent = Inches(-0.25)
        run = p.add_run(f"[{i}] ")
        run.font.size = Pt(size)
        run2 = p.add_run(ref.raw)
        run2.font.size = Pt(size)


# --- per-publisher renderers ------------------------------------------------

def _render_ieee_docx(paper: Paper):
    doc = Document()
    _base_font(doc, "Times New Roman", 10)
    sec = doc.sections[0]
    sec.page_height, sec.page_width = Inches(11), Inches(8.5)
    sec.top_margin = sec.bottom_margin = Inches(0.75)
    sec.left_margin = sec.right_margin = Inches(0.625)

    # Title block spans both columns (single-column section first).
    _para(doc, paper.title, size=24, bold=False,
          align=WD_ALIGN_PARAGRAPH.CENTER, space_after=8, font="Times New Roman")
    if paper.authors:
        names = ",  ".join(a.name for a in paper.authors)
        _para(doc, names, size=11, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
        affils = [a.affiliation for a in paper.authors if a.affiliation]
        for af in dict.fromkeys(affils):
            _para(doc, af, size=10, italic=True,
                  align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)

    # Switch to two columns for the body.
    body = doc.add_section(WD_SECTION.CONTINUOUS)
    body.top_margin = body.bottom_margin = Inches(0.75)
    body.left_margin = body.right_margin = Inches(0.625)
    _set_columns(body, 2)

    if paper.abstract:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(6)
        r = p.add_run("Abstract—")
        r.bold = True
        r.italic = True
        r.font.size = Pt(9)
        r2 = p.add_run(paper.abstract)
        r2.italic = True
        r2.font.size = Pt(9)
    if paper.keywords:
        p = doc.add_paragraph()
        r = p.add_run("Index Terms—")
        r.bold = True
        r.italic = True
        r.font.size = Pt(9)
        r2 = p.add_run(", ".join(paper.keywords))
        r2.italic = True
        r2.font.size = Pt(9)

    _write_sections(doc, paper.sections, 10, numbering="roman")
    _write_references(doc, paper, 9)
    return doc


def _render_springer_docx(paper: Paper):
    doc = Document()
    _base_font(doc, "Times New Roman", 10)
    sec = doc.sections[0]
    sec.top_margin = sec.bottom_margin = Inches(1.0)
    sec.left_margin = sec.right_margin = Inches(1.2)

    _para(doc, paper.title, size=16, bold=True,
          align=WD_ALIGN_PARAGRAPH.LEFT, space_after=8)
    if paper.authors:
        _para(doc, " and ".join(a.name for a in paper.authors), size=11,
              space_after=2)
        affils = [a.affiliation for a in paper.authors if a.affiliation]
        for af in dict.fromkeys(affils):
            _para(doc, af, size=9, italic=True, space_after=8)

    if paper.abstract:
        _para(doc, "Abstract.", size=9, bold=True, space_after=0)
        p = _para(doc, paper.abstract, size=9, align=WD_ALIGN_PARAGRAPH.JUSTIFY,
                  space_after=4)
        p.paragraph_format.left_indent = Inches(0.4)
        p.paragraph_format.right_indent = Inches(0.4)
    if paper.keywords:
        p = _para(doc, "Keywords: " + " · ".join(paper.keywords), size=9,
                  bold=False, space_after=10)
        p.paragraph_format.left_indent = Inches(0.4)
        p.paragraph_format.right_indent = Inches(0.4)

    _write_sections(doc, paper.sections, 10, numbering="arabic")
    _write_references(doc, paper, 9)
    return doc


def _render_elsevier_docx(paper: Paper):
    doc = Document()
    _base_font(doc, "Times New Roman", 12)
    sec = doc.sections[0]
    sec.top_margin = sec.bottom_margin = Inches(1.0)
    sec.left_margin = sec.right_margin = Inches(1.0)

    _para(doc, paper.title, size=18, bold=True,
          align=WD_ALIGN_PARAGRAPH.LEFT, space_after=10)
    for a in paper.authors:
        _para(doc, a.name, size=12, space_after=0)
        if a.affiliation:
            _para(doc, a.affiliation, size=10, italic=True, space_after=2)
        if a.email:
            _para(doc, a.email, size=10, italic=True, space_after=6)

    if paper.abstract:
        _para(doc, "Abstract", size=12, bold=True, space_after=2)
        _para(doc, paper.abstract, size=11, align=WD_ALIGN_PARAGRAPH.JUSTIFY,
              space_after=6)
    if paper.keywords:
        _para(doc, "Keywords: " + ", ".join(paper.keywords), size=11,
              space_after=10)

    _write_sections(doc, paper.sections, 12, numbering="arabic")
    _write_references(doc, paper, 11)
    return doc


_RENDERERS = {
    "ieee": _render_ieee_docx,
    "springer": _render_springer_docx,
    "elsevier": _render_elsevier_docx,
}


def render_docx(paper: Paper, publisher: str, out_path: str):
    publisher = publisher.lower()
    if publisher not in _RENDERERS:
        raise ValueError(f"Unknown publisher '{publisher}'. "
                         f"Choose from {sorted(_RENDERERS)}.")
    doc = _RENDERERS[publisher](paper)
    doc.save(out_path)
    return out_path
