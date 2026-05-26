"""Humanize / naturalize text: rewrite stiff, AI-sounding prose into plainer
language.

This is a deterministic, offline rewriter. It does three kinds of edits:
  * swap inflated vocabulary for plain words ("utilize" -> "use"),
  * drop empty filler phrases ("it is important to note that"),
  * optionally use contractions and split overlong sentences.

It is a writing aid for improving your *own* drafts' readability. It is not an
"AI-detector bypass" and makes no such guarantee. For higher-quality rewriting,
set an LLM backend via ``set_llm_backend`` (see the optional hook at the bottom).
"""

from __future__ import annotations

import re
from dataclasses import dataclass, field
from pathlib import Path
from typing import Callable, List, Optional

LEVELS = ["light", "medium", "strong"]


# --- rewrite tables ---------------------------------------------------------

# Multi-word clichés and inflated phrases -> plainer wording. An empty value
# removes the phrase entirely (sentence is re-capitalized afterwards).
PHRASE_MAP = [
    ("it is important to note that", ""),
    ("it is worth noting that", ""),
    ("it should be noted that", ""),
    ("it is worth mentioning that", ""),
    ("needless to say,", ""),
    ("as a matter of fact,", ""),
    ("at the end of the day,", ""),
    ("when it comes to", "for"),
    ("in order to", "to"),
    ("due to the fact that", "because"),
    ("owing to the fact that", "because"),
    ("in light of the fact that", "because"),
    ("in the event that", "if"),
    ("in the event of", "if there is"),
    ("for the purpose of", "to"),
    ("with the aim of", "to"),
    ("a large number of", "many"),
    ("a significant number of", "many"),
    ("a wide range of", "many"),
    ("a plethora of", "many"),
    ("a myriad of", "many"),
    ("a majority of", "most"),
    ("the vast majority of", "most"),
    ("a small number of", "a few"),
    ("plays a crucial role in", "is central to"),
    ("plays a vital role in", "is central to"),
    ("plays a key role in", "is central to"),
    ("plays a significant role in", "matters in"),
    ("in the realm of", "in"),
    ("in the world of", "in"),
    ("with regard to", "about"),
    ("with respect to", "about"),
    ("in terms of", "for"),
    ("prior to", "before"),
    ("subsequent to", "after"),
    ("in conclusion,", ""),
    ("in summary,", ""),
    ("to summarize,", ""),
    ("first and foremost,", "first,"),
    ("last but not least,", "finally,"),
    ("a testament to", "proof of"),
    ("stands as a testament to", "shows"),
    ("paving the way for", "enabling"),
    ("paves the way for", "enables"),
    ("the ever-evolving", "the changing"),
    ("ever-evolving", "changing"),
    ("an integral part of", "a key part of"),
    ("in today's fast-paced world,", ""),
    ("in today's digital age,", ""),
    ("it goes without saying that", ""),
]

# Inflated single words -> plain ones. Forms are listed explicitly so we keep
# the right tense/number. Matching is whole-word, case-insensitive; the
# replacement copies the original word's capitalization.
WORD_MAP = [
    ("utilize", "use"), ("utilizes", "uses"), ("utilized", "used"),
    ("utilizing", "using"), ("utilization", "use"),
    ("leverage", "use"), ("leverages", "uses"), ("leveraging", "using"),
    ("facilitate", "help"), ("facilitates", "helps"), ("facilitated", "helped"),
    ("facilitating", "helping"),
    ("delve", "look"), ("delves", "looks"), ("delving", "looking"),
    ("commence", "start"), ("commences", "starts"), ("commenced", "started"),
    ("endeavor", "try"), ("endeavors", "tries"),
    ("ascertain", "find out"), ("demonstrate", "show"),
    ("demonstrates", "shows"), ("demonstrated", "showed"),
    ("subsequently", "then"), ("furthermore", "also"), ("moreover", "also"),
    ("additionally", "also"), ("consequently", "so"), ("therefore", "so"),
    ("nevertheless", "still"), ("nonetheless", "still"),
    ("approximately", "about"), ("numerous", "many"), ("myriad", "many"),
    ("plethora", "plenty"), ("pivotal", "key"), ("crucial", "key"),
    ("paramount", "key"), ("robust", "strong"), ("seamless", "smooth"),
    ("seamlessly", "smoothly"), ("comprehensive", "complete"),
    ("intricate", "complex"), ("multifaceted", "complex"),
    ("realm", "area"), ("landscape", "area"), ("underscore", "highlight"),
    ("underscores", "highlights"), ("underscored", "highlighted"),
    ("showcase", "show"), ("showcases", "shows"), ("showcasing", "showing"),
    ("encompass", "include"), ("encompasses", "includes"),
    ("elucidate", "explain"), ("elucidates", "explains"),
    ("ameliorate", "improve"), ("optimal", "best"),
    ("aforementioned", "above"), ("henceforth", "from now on"),
]

# Filler / hedge words removed at the 'strong' level (or when toggled on).
FILLER_WORDS = [
    "very", "really", "quite", "actually", "basically", "essentially",
    "literally", "simply", "just", "definitely", "certainly", "arguably",
    "notably", "interestingly", "importantly",
]

CONTRACTIONS = [
    ("do not", "don't"), ("does not", "doesn't"), ("did not", "didn't"),
    ("is not", "isn't"), ("are not", "aren't"), ("was not", "wasn't"),
    ("were not", "weren't"), ("has not", "hasn't"), ("have not", "haven't"),
    ("had not", "hadn't"), ("will not", "won't"), ("would not", "wouldn't"),
    ("can not", "can't"), ("cannot", "can't"), ("could not", "couldn't"),
    ("should not", "shouldn't"), ("it is", "it's"), ("that is", "that's"),
    ("there is", "there's"), ("they are", "they're"), ("we are", "we're"),
    ("you are", "you're"), ("we will", "we'll"), ("they will", "they'll"),
    ("we have", "we've"), ("they have", "they've"),
]


# --- options ----------------------------------------------------------------

@dataclass
class Options:
    simplify_vocab: bool = True      # apply WORD_MAP + PHRASE_MAP
    remove_filler: bool = False      # drop FILLER_WORDS
    use_contractions: bool = False   # apply CONTRACTIONS
    split_sentences: bool = False    # break very long sentences
    max_sentence_words: int = 32     # threshold for splitting

    @classmethod
    def for_level(cls, level: str) -> "Options":
        level = (level or "medium").lower()
        if level == "light":
            return cls(simplify_vocab=True)
        if level == "strong":
            return cls(simplify_vocab=True, remove_filler=True,
                       use_contractions=True, split_sentences=True)
        return cls(simplify_vocab=True, remove_filler=True,
                   split_sentences=True)  # medium


@dataclass
class Result:
    original: str
    text: str
    changes: int = 0
    notes: List[str] = field(default_factory=list)


# --- helpers ----------------------------------------------------------------

def _match_case(original: str, replacement: str) -> str:
    if original[:1].isupper():
        return replacement[:1].upper() + replacement[1:]
    return replacement


def _apply_phrases(text: str, counter: List[int]) -> str:
    for phrase, repl in PHRASE_MAP:
        regex = re.compile(re.escape(phrase).replace(r"\ ", r"\s+"), re.IGNORECASE)

        def sub(m, repl=repl):
            counter[0] += 1
            return _match_case(m.group(0), repl) if repl else ""

        text = regex.sub(sub, text)
    return text


def _apply_words(text: str, mapping, counter: List[int]) -> str:
    for word, repl in mapping:
        regex = re.compile(r"\b" + re.escape(word) + r"\b", re.IGNORECASE)

        def sub(m, repl=repl):
            counter[0] += 1
            return _match_case(m.group(0), repl)

        text = regex.sub(sub, text)
    return text


def _remove_filler(text: str, counter: List[int]) -> str:
    for word in FILLER_WORDS:
        # Remove the word plus one adjoining space; keep sentence intact.
        regex = re.compile(r"\b" + re.escape(word) + r"\b\s?", re.IGNORECASE)
        text, n = regex.subn("", text)
        counter[0] += n
    return text


def _split_long_sentences(text: str, max_words: int, counter: List[int]) -> str:
    # Easiest safe split: turn semicolons joining clauses into separate
    # sentences when the sentence is long.
    sentences = re.split(r"(?<=[.!?])\s+", text)
    out = []
    for s in sentences:
        if len(s.split()) > max_words and "; " in s:
            parts = [p.strip() for p in s.split("; ") if p.strip()]
            parts = [p if p.endswith((".", "!", "?")) else p + "." for p in parts]
            parts = [p[:1].upper() + p[1:] for p in parts]
            counter[0] += len(parts) - 1
            out.append(" ".join(parts))
        else:
            out.append(s)
    return " ".join(out)


def _cleanup(text: str) -> str:
    text = re.sub(r"[ \t]{2,}", " ", text)
    text = re.sub(r"\s+([,.;:!?])", r"\1", text)   # space before punctuation
    text = re.sub(r"\(\s+", "(", text)
    text = re.sub(r"\s+\)", ")", text)
    text = re.sub(r",\s*,", ",", text)
    text = text.strip()
    # Re-capitalize sentence starts (helps after leading phrase removals).
    def cap(m):
        return m.group(1) + m.group(2).upper()
    text = re.sub(r"(^\s*|[.!?]\s+)([a-z])", cap, text)
    return text.strip()


# --- public API -------------------------------------------------------------

_LLM_BACKEND: Optional[Callable[[str, str], str]] = None


def set_llm_backend(fn: Optional[Callable[[str, str], str]]) -> None:
    """Register an optional LLM rewriter: fn(text, level) -> rewritten text.

    When set, :func:`humanize_text` uses it instead of the rule engine. This is
    the hook for plugging in a higher-quality model later (e.g. the Anthropic
    SDK). Pass ``None`` to go back to the offline rule engine.
    """
    global _LLM_BACKEND
    _LLM_BACKEND = fn


def humanize_text(text: str, level: str = "medium",
                  options: Optional[Options] = None) -> Result:
    """Rewrite ``text`` into plainer language. Returns a :class:`Result`."""
    if text is None:
        text = ""
    original = text
    if _LLM_BACKEND is not None:
        rewritten = _LLM_BACKEND(text, level)
        return Result(original=original, text=rewritten, changes=-1,
                      notes=["Rewritten by the configured LLM backend."])

    opts = options or Options.for_level(level)
    counter = [0]
    if opts.simplify_vocab:
        text = _apply_phrases(text, counter)
        text = _apply_words(text, WORD_MAP, counter)
    if opts.use_contractions:
        text = _apply_words(text, CONTRACTIONS, counter)
    if opts.remove_filler:
        text = _remove_filler(text, counter)
    if opts.split_sentences:
        text = _split_long_sentences(text, opts.max_sentence_words, counter)
    text = _cleanup(text)
    return Result(original=original, text=text, changes=counter[0])


# --- file-level helpers -----------------------------------------------------

def extract_paragraphs(path: str) -> List[str]:
    """Return the body paragraphs of a file (abstract + section text)."""
    from .parsers import parse
    paper = parse(path)
    paras: List[str] = []
    if paper.abstract:
        paras.append(paper.abstract)

    def walk(sections):
        for sec in sections:
            for p in sec.paragraphs:
                paras.append(p)
            walk(sec.children)

    walk(paper.sections)
    return paras


def humanize_paragraphs(paragraphs: List[str], scope="all",
                        level: str = "medium",
                        options: Optional[Options] = None) -> List[Result]:
    """Humanize selected paragraphs.

    ``scope`` is "all", a single 0-based index, or a list of indices.
    Paragraphs not selected are returned unchanged (changes=0).
    """
    if scope == "all":
        selected = set(range(len(paragraphs)))
    elif isinstance(scope, int):
        selected = {scope}
    else:
        selected = set(scope)

    results = []
    for i, para in enumerate(paragraphs):
        if i in selected:
            results.append(humanize_text(para, level, options))
        else:
            results.append(Result(original=para, text=para, changes=0))
    return results


def humanize_file(path: str, scope="all", level: str = "medium",
                  outdir: str = ".", options: Optional[Options] = None):
    """Humanize a whole file (or selected paragraphs) and write outputs.

    Produces ``<stem>_humanized.txt`` (and ``.docx``) in ``outdir``. Returns
    ``(output_paths, results)``.
    """
    paragraphs = extract_paragraphs(path)
    results = humanize_paragraphs(paragraphs, scope, level, options)

    out_dir = Path(outdir)
    out_dir.mkdir(parents=True, exist_ok=True)
    stem = Path(path).stem
    txt_path = out_dir / f"{stem}_humanized.txt"
    body = "\n\n".join(r.text for r in results)
    txt_path.write_text(body, encoding="utf-8")

    outputs = [str(txt_path)]
    try:
        from docx import Document
        doc = Document()
        for r in results:
            doc.add_paragraph(r.text)
        docx_path = out_dir / f"{stem}_humanized.docx"
        doc.save(str(docx_path))
        outputs.append(str(docx_path))
    except Exception:
        pass

    return outputs, results
